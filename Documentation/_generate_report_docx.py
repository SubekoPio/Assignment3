"""
Generate EduManage System Report as a styled Word (.docx) document.
Run once to produce the output file, then this script can be deleted.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Branding colours ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0F, 0x4C, 0x81)
CORAL  = RGBColor(0xF2, 0x8C, 0x6F)
SLATE  = RGBColor(0x33, 0x41, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xEE, 0xF3, 0xF9)
BLACK  = RGBColor(0x1F, 0x29, 0x37)
GRAY   = RGBColor(0x64, 0x74, 0x8B)

GROUP_MEMBERS = [
    ("Member Name 1", "REG/XXXX/001"),
    ("Member Name 2", "REG/XXXX/002"),
    ("Member Name 3", "REG/XXXX/003"),
    ("Member Name 4", "REG/XXXX/004"),
    ("Member Name 5", "REG/XXXX/005"),
]

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=NAVY, space_before=14, space_after=6):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p


def add_body(doc, text, space_before=2, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = BLACK
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.5 * (level + 1))
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    for run in p.runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = BLACK
    return p


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent  = Cm(0.5 * (level + 1))
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = BLACK
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(10)
    run.font.color.rgb = SLATE
    return p


def add_divider(doc, color_hex="0F4C81"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p


# ── Document build ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Default paragraph font
doc.styles["Normal"].font.name  = "Calibri"
doc.styles["Normal"].font.size  = Pt(11)
doc.styles["Normal"].font.color.rgb = BLACK


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Coloured title banner table
banner = doc.add_table(rows=1, cols=1)
banner.style = "Table Grid"
cell = banner.cell(0, 0)
set_cell_bg(cell, "0F4C81")
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(14)
run = p.add_run("EDUMANAGE SYSTEM REPORT")
run.font.name  = "Calibri"
run.font.size  = Pt(22)
run.font.bold  = True
run.font.color.rgb = WHITE

doc.add_paragraph()

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EduManage — Advanced Education Management System")
run.font.name  = "Calibri"
run.font.size  = Pt(14)
run.font.color.rgb = NAVY
run.font.bold  = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Intermediate Education System Project")
run2.font.size = Pt(12)
run2.font.color.rgb = SLATE

doc.add_paragraph()
add_divider(doc, "F28C6F")
doc.add_paragraph()

# Group members table
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GROUP MEMBERS")
run.font.name  = "Calibri"
run.font.size  = Pt(13)
run.font.bold  = True
run.font.color.rgb = NAVY

doc.add_paragraph()

members_table = doc.add_table(rows=len(GROUP_MEMBERS) + 1, cols=2)
members_table.style = "Table Grid"
members_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Header row
hcells = members_table.rows[0].cells
hcells[0].text = "Full Name"
hcells[1].text = "Registration Number"
for i, cell in enumerate(hcells):
    set_cell_bg(cell, "0F4C81")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(11)

# Member rows
for idx, (name, reg) in enumerate(GROUP_MEMBERS):
    row = members_table.rows[idx + 1]
    row.cells[0].text = name
    row.cells[1].text = reg
    bg = "EEF3F9" if idx % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        set_cell_bg(cell, bg)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK

doc.add_paragraph()
add_divider(doc, "0F4C81")

# Date / submission line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Date: June 2026   |   Academic Year: 2025/2026")
run.font.size  = Pt(11)
run.font.color.rgb = GRAY
run.font.italic = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Abstract", level=1)
add_divider(doc)
add_body(doc,
    "EduManage is a desktop-based education management system developed to centralize and "
    "streamline academic administration. The system manages students, teachers, courses, units, "
    "enrollments, and grades while providing reporting and analytics capabilities. This report "
    "presents the problem context, project objectives, system architecture, implementation details, "
    "operational workflow, and testing/evaluation outcomes. The final system uses a layered design, "
    "CSV persistence, and a feature-complete graphical interface with validated business rules and "
    "automated test coverage."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 1: Problem Description and Justification", level=1)
add_divider(doc)

add_heading(doc, "1.1 Background", level=2)
add_body(doc,
    "Educational organizations frequently rely on fragmented record-keeping methods such as "
    "spreadsheets, manual registers, and disconnected digital files. These practices increase "
    "operational risk and reduce administrative efficiency."
)

add_heading(doc, "1.2 Problem Statement", level=2)
add_body(doc, "The identified problems are:")
for item in [
    "Record inconsistency across multiple data sources.",
    "Duplicate data entry and validation gaps.",
    "Difficulty in managing unit-level course progression.",
    "Limited visibility of teacher workload and grade distributions.",
    "Delayed transcript/report preparation and analytics reporting.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.3 Justification", level=2)
add_body(doc, "EduManage is justified as a centralized, maintainable, and practical desktop platform that:")
for item in [
    "Enforces data integrity through business-rule validation.",
    "Provides end-to-end academic workflows in one interface.",
    "Supports professional report export and decision-support analytics.",
    "Preserves data in transparent, portable CSV format.",
]:
    add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 2: Objectives of the System", level=1)
add_divider(doc)

add_heading(doc, "2.1 General Objective", level=2)
add_body(doc,
    "Develop a reliable education-management information system that consolidates student, "
    "teacher, course, enrollment, and reporting workflows."
)

add_heading(doc, "2.2 Specific Objectives", level=2)
for item in [
    "Implement full CRUD operations for core entities (students, courses, teachers).",
    "Model unit-based course structures and enforce globally unique unit IDs.",
    "Support teacher assignment at both course and unit levels.",
    "Enable unit-level enrollment and grading workflows.",
    "Compute weighted course GPA and overall CGPA automatically.",
    "Provide polished report preview and professional PDF export.",
    "Provide analytics for enrollment, grades, and workload.",
    "Validate functionality using positive and negative automated test suites.",
]:
    add_numbered(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 3: System Design", level=1)
add_divider(doc)

add_heading(doc, "3.1 Architectural Overview", level=2)
add_body(doc,
    "EduManage follows a four-layer architecture. Each layer has a single responsibility "
    "and interacts only with adjacent layers:"
)

arch_table = doc.add_table(rows=4, cols=2)
arch_table.style = "Table Grid"
layers = [
    ("Presentation Layer", "CustomTkinter GUI — handles all user interaction and display."),
    ("Application Logic Layer", "GUI event handlers — translates user actions into service calls."),
    ("Business Logic Layer", "EducationSystem (system.py) — enforces rules, computes results."),
    ("Persistence Layer", "CSV + JSON column files — durable storage and reconstruction."),
]
for i, (layer, desc) in enumerate(layers):
    cells = arch_table.rows[i].cells
    set_cell_bg(cells[0], "0F4C81" if i % 2 == 0 else "EEF3F9")
    cells[0].text = layer
    cells[1].text = desc
    for run in cells[0].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = WHITE if i % 2 == 0 else NAVY
        run.font.size  = Pt(11)
    for run in cells[1].paragraphs[0].runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = BLACK

doc.add_paragraph()

add_heading(doc, "3.2 Design Rationale", level=2)
for item in [
    "GUI handles interaction and presentation only.",
    "EducationSystem encapsulates business rules and operations.",
    "models.py captures domain entities and reusable behavior.",
    "CSV files provide persistent storage with simple portability.",
]:
    add_bullet(doc, item)

add_heading(doc, "3.3 Structural Components", level=2)
components = [
    ("gui_main.py", "Theme handling, tab creation, form interaction, data display, report generation, PDF export, Matplotlib rendering."),
    ("system.py", "Service-layer methods for CRUD, assignment, enrollment, grading, analytics, and persistence."),
    ("models.py", "Class definitions and entity-level methods (Person, Student, Teacher, Course, Unit)."),
    ("Data_Storage(CSV)", "Runtime durable CSV data files."),
    ("tests/", "Automated system and negative-path verification (26 tests total)."),
]
comp_table = doc.add_table(rows=len(components), cols=2)
comp_table.style = "Table Grid"
for i, (comp, desc) in enumerate(components):
    cells = comp_table.rows[i].cells
    set_cell_bg(cells[0], "EEF3F9")
    cells[0].text = comp
    cells[1].text = desc
    for run in cells[0].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = NAVY
        run.font.name  = "Courier New"
        run.font.size  = Pt(10)
    for run in cells[1].paragraphs[0].runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = BLACK

doc.add_paragraph()

add_heading(doc, "3.4 Core Process Flows", level=2)

add_heading(doc, "3.4.1 Enrollment and Grading Flow", level=3)
flow_steps = [
    "Select Student / Course / Unit from GUI controls.",
    "Validate references and state (EducationSystem).",
    "Enroll unit.",
    "Assign unit grade.",
    "Map grade → letter → GPA points.",
    "Build report data and analytics.",
    "Save to CSV.",
]
for step in flow_steps:
    add_numbered(doc, step)

add_heading(doc, "3.4.2 Reporting Flow", level=3)
report_flow = [
    "Choose student from dropdown.",
    "Collect enrolled courses and units from student record.",
    "Compute weighted GPA per course and overall CGPA.",
    "Render preview text in GUI textbox.",
    "Export styled PDF via ReportLab (branded layout, grade badges, signature, footer).",
]
for step in report_flow:
    add_numbered(doc, step)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 4: Implementation Details", level=1)
add_divider(doc)

add_heading(doc, "4.1 Technology Stack", level=2)
tech = [
    ("Python", "Core implementation language."),
    ("CustomTkinter / tkinter / ttk", "UI components and table views."),
    ("Matplotlib", "Embedded analytics charts."),
    ("ReportLab", "Professional PDF report export."),
    ("Pytest", "Automated testing."),
    ("CSV + JSON", "Persistent data representation."),
    ("python-docx", "Word document generation."),
]
tech_table = doc.add_table(rows=len(tech), cols=2)
tech_table.style = "Table Grid"
for i, (tech_name, desc) in enumerate(tech):
    cells = tech_table.rows[i].cells
    set_cell_bg(cells[0], "0F4C81" if i % 2 == 0 else "EEF3F9")
    cells[0].text = tech_name
    cells[1].text = desc
    for run in cells[0].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = WHITE if i % 2 == 0 else NAVY
        run.font.size  = Pt(11)
    for run in cells[1].paragraphs[0].runs:
        run.font.size  = Pt(11)

doc.add_paragraph()

add_heading(doc, "4.2 Domain Model Implementation", level=2)
models = [
    ("Person", "Base class with person_id, name, email. Enforces email format via regex."),
    ("Student", "Stores unit-level enrollment dict by course. Supports enrollment and grade assignment."),
    ("Teacher", "Stores department, assigned_courses, and taught_units mappings."),
    ("Course", "Stores metadata, teacher_id, teacher_ids list, and units list."),
    ("Unit", "Represents individual curriculum component with credits."),
]
model_table = doc.add_table(rows=len(models), cols=2)
model_table.style = "Table Grid"
for i, (model, desc) in enumerate(models):
    cells = model_table.rows[i].cells
    set_cell_bg(cells[0], "EEF3F9")
    cells[0].text = model
    cells[1].text = desc
    for run in cells[0].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = NAVY
        run.font.size  = Pt(11)
    for run in cells[1].paragraphs[0].runs:
        run.font.size  = Pt(11)

doc.add_paragraph()

add_heading(doc, "4.3 Service Layer (EducationSystem) Implementation", level=2)
add_body(doc, "Implemented method groups include:")
for item in [
    "ID generation: next_student_id(), next_teacher_id(), next_course_id(), next_unit_id().",
    "CRUD operations for students, courses, teachers, and units.",
    "Teacher assignment: assign_teacher_to_course(), assign_teacher_to_unit().",
    "Enrollment: enroll_student_unit(), remove_unit_enrollment().",
    "Grading: assign_unit_grade().",
    "Report generation: get_student_report().",
    "Analytics: get_analytics(), _get_students_per_course(), _get_grades_distribution(), _get_teacher_workload().",
    "Grade mapping: grade_to_letter(), grade_to_point().",
    "Persistence: save_data(), load_data(), create_backup_zip().",
    "CSV export: export_courses_summary().",
]:
    add_bullet(doc, item)

add_heading(doc, "4.4 GUI Implementation", level=2)
add_body(doc, "The interface is organized into six tabs:")
for tab in [
    "Students — CRUD, search, selection.",
    "Courses — CRUD, units panel, manage-units dialog.",
    "Teachers — CRUD, course/unit assignment.",
    "Enrollment & Grades — unit enrollment, grade assignment.",
    "Reports — preview generation, PDF export.",
    "Analysis — analytics dashboard with snapshots.",
]:
    add_bullet(doc, tab)

add_heading(doc, "4.5 Data Persistence and Reconstruction", level=2)
for item in [
    "Students, courses, teachers, and enrollments are written to dedicated CSV files.",
    "Course units and teacher ID lists are serialized as JSON inside CSV columns.",
    "On load, object graphs and cross-links are fully reconstructed.",
    "Unit-teacher relationships and student unit-grade records are restored from CSV.",
]:
    add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 5: System Functionality and Features", level=1)
add_divider(doc)

features = {
    "5.1 Student Features": [
        "Create, read, update, delete students.",
        "Auto-ID support and email validation.",
        "Search/filter student list in real time.",
    ],
    "5.2 Course and Unit Features": [
        "Create, update, delete courses.",
        "Add, update, delete units under courses.",
        "Globally unique unit ID enforcement.",
        "Display selected course units and assigned teachers.",
    ],
    "5.3 Teacher Features": [
        "Create, update, delete teachers.",
        "Assign teacher to whole courses.",
        "Assign teacher to specific units.",
        "Track teacher workload from taught units.",
    ],
    "5.4 Enrollment and Grading Features": [
        "Enroll students at unit level.",
        "Assign and update grades by selected enrollment.",
        "Remove enrollment records safely.",
        "Enforce dependency checks before grade assignment.",
    ],
    "5.5 Reporting Features": [
        "Generate report preview with student, course, and unit details.",
        "Compute and show course GPA and overall CGPA.",
        "Export branded PDF: header, color table, grade badges, signatures, footer.",
        "Optional institution logo auto-detection.",
    ],
    "5.6 Analytics Features": [
        "Students per course bar chart.",
        "Grade distribution by letter chart.",
        "Teacher workload horizontal bar chart.",
        "Summary statistics panel.",
        "Student/teacher snapshot via dropdown selections.",
    ],
}

for section_title, items in features.items():
    add_heading(doc, section_title, level=2)
    for item in items:
        add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 6: How the System Works (Input, Processing, Output)", level=1)
add_divider(doc)

add_heading(doc, "6.1 Input", level=2)
add_body(doc, "Input sources include:")
for item in [
    "Text fields for names, emails, departments, credits, grades.",
    "Dropdown selectors for students, courses, units, teachers.",
    "Tree/list selections for record targeting.",
    "Action buttons for CRUD, assignment, enrollment, reporting, analytics.",
]:
    add_bullet(doc, item)

add_heading(doc, "6.2 Processing", level=2)
add_body(doc, "Processing flow in EducationSystem:")
for item in [
    "Validate identity references and rule constraints.",
    "Execute domain operation.",
    "Update in-memory entity objects and relationships.",
    "Compute derived values: letters, points, GPA, CGPA, analytics metrics.",
    "Persist updated state to CSV files.",
]:
    add_numbered(doc, item)

add_heading(doc, "6.3 Output", level=2)
add_body(doc, "Outputs include:")
for item in [
    "Updated tables and forms in GUI.",
    "Report preview text in textbox.",
    "Styled PDF report file exported to selected path.",
    "Matplotlib analytics charts.",
    "Updated CSV persistent state.",
]:
    add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 7: Testing and Evaluation", level=1)
add_divider(doc)

add_heading(doc, "7.1 Testing Strategy", level=2)
add_body(doc, "Two complementary suites were used:")
for item in [
    "test_system.py — positive end-to-end workflows (11 tests).",
    "test_system_negative.py — invalid/exception workflows (15 tests).",
]:
    add_bullet(doc, item)

add_heading(doc, "7.2 Executed Test Results", level=2)

results_table = doc.add_table(rows=4, cols=2)
results_table.style = "Table Grid"
results_data = [
    ("Test Suite", "Result"),
    ("test_system.py — Functional Suite", "11 / 11 PASSED"),
    ("test_system_negative.py — Negative Suite", "15 / 15 PASSED"),
    ("Combined Total", "26 / 26 PASSED"),
]
for i, (left, right) in enumerate(results_data):
    cells = results_table.rows[i].cells
    bg = "0F4C81" if i == 0 else ("EEF3F9" if i % 2 == 0 else "FFFFFF")
    fgL = WHITE if i == 0 else NAVY
    fgR = WHITE if i == 0 else (GRAY if i < 3 else NAVY)
    set_cell_bg(cells[0], bg)
    set_cell_bg(cells[1], bg)
    cells[0].text = left
    cells[1].text = right
    for run in cells[0].paragraphs[0].runs:
        run.font.bold  = (i == 0 or i == 3)
        run.font.color.rgb = fgL
        run.font.size  = Pt(11)
    for run in cells[1].paragraphs[0].runs:
        run.font.bold  = (i == 0 or i == 3)
        run.font.color.rgb = fgR
        run.font.size  = Pt(11)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_heading(doc, "7.3 Evaluated Quality Dimensions", level=2)
for item in [
    "Correctness of domain rules.",
    "Consistency of GPA and CGPA mathematics.",
    "Cascade behavior on delete operations.",
    "Persistence round-trip reliability.",
    "Robustness against invalid actions.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.4 Limitations and Future Work", level=2)
add_body(doc, "Current limitations:")
for item in [
    "CSV storage is suitable for desktop/single-user scope, not concurrent multi-user scale.",
    "Security/authentication is not currently implemented.",
]:
    add_bullet(doc, item)

add_body(doc, "Recommended improvements:")
for item in [
    "Migrate to a transactional database backend (SQLite/PostgreSQL).",
    "Add automated GUI interaction test layer.",
    "Integrate CI coverage thresholds.",
    "Add role-based access and audit log support.",
]:
    add_numbered(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Chapter 8: Conclusion", level=1)
add_divider(doc)
add_body(doc,
    "EduManage successfully fulfills the defined academic-management requirements through a "
    "layered architecture, robust service logic, feature-rich GUI workflows, professional "
    "reporting, and validated test outcomes. The system is maintainable, demonstrable, and "
    "extensible, with clear pathways for scaling and hardening in future iterations. "
    "The project demonstrates practical application of object-oriented design, clean "
    "architecture, and professional-grade software engineering practices."
)

doc.add_paragraph()
add_divider(doc, "F28C6F")

# Final sign-off
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("End of Report  —  EduManage Group Project  —  June 2026")
run.font.size   = Pt(10)
run.font.color.rgb = GRAY
run.font.italic = True


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "EduManage_System_Report.docx")
doc.save(out_path)
print(f"✓ Word document saved to:\n  {out_path}")
